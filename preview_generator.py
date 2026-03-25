"""
Preview generator module for Legal AI Service.
Generates preview documents with blur effect for unpaid documents.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def estimate_total_pages(text_content):
    """
    Estimate number of pages based on text length.
    Rough estimate: ~3000 characters per page (with formatting).
    """
    if not text_content:
        return 1
    chars_per_page = 3000
    return max(1, int(len(text_content) / chars_per_page))


def split_content_for_preview(content, visible_ratio=0.25):
    """
    Split content into visible and blurred parts.
    
    Args:
        content: Full document content
        visible_ratio: How much to show (0.25 = 25%)
    
    Returns:
        tuple: (visible_content, is_truncated)
    """
    if not content:
        return content, False
    
    lines = content.strip().split('\n')
    total_lines = len(lines)
    
    if total_lines <= 5:
        # For very short documents, show first 2 lines
        visible_lines = min(2, total_lines)
        visible = '\n'.join(lines[:visible_lines])
        return visible, total_lines > visible_lines
    
    # Calculate visible lines (at least 5, at most 30% of total)
    visible_count = max(5, int(total_lines * visible_ratio))
    visible_count = min(visible_count, total_lines)
    
    visible_content = '\n'.join(lines[:visible_count])
    
    return visible_content, total_lines > visible_count


def generate_preview_html(case_data, visible_content, document_type):
    """
    Generate HTML preview with blur effect.
    
    Args:
        case_data: Case information
        visible_content: Visible portion of content
        document_type: Type of document
    
    Returns:
        HTML string
    """
    doc_type_names = {
        'complaint': 'Исковое заявление',
        'appeal': 'Апелляционная жалоба',
        'objection': 'Возражение',
        'strategy': 'Стратегия защиты',
        'petition': 'Ходатайство',
        'contract': 'Договор',
        'claim': 'Претензия'
    }
    
    doc_type = doc_type_names.get(document_type, 'Юридический документ')
    
    # Format visible content
    formatted_content = format_content_to_html(visible_content)
    
    html = f'''<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Предпросмотр - {doc_type}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Times New Roman', Times, serif;
            background: #f5f5f5;
            padding: 20px;
            line-height: 1.6;
        }}
        
        .preview-container {{
            max-width: 800px;
            margin: 0 auto;
            background: white;
            padding: 60px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            position: relative;
        }}
        
        .document-header {{
            text-align: center;
            margin-bottom: 30px;
            padding-bottom: 20px;
            border-bottom: 2px solid #333;
        }}
        
        .document-title {{
            font-size: 18px;
            font-weight: bold;
            text-transform: uppercase;
            margin-bottom: 10px;
        }}
        
        .document-type {{
            font-size: 14px;
            color: #666;
        }}
        
        .content {{
            font-size: 14px;
            text-align: justify;
        }}
        
        .content p {{
            margin-bottom: 12px;
            text-indent: 20px;
        }}
        
        .content h1, .content h2, .content h3 {{
            margin: 20px 0 10px 0;
            font-weight: bold;
        }}
        
        .content h1 {{
            font-size: 16px;
            text-align: center;
        }}
        
        .content h2 {{
            font-size: 14px;
        }}
        
        .blur-overlay {{
            position: relative;
            height: 400px;
            margin-top: 20px;
            overflow: hidden;
        }}
        
        .blur-content {{
            filter: blur(8px);
            user-select: none;
            pointer-events: none;
            opacity: 0.3;
        }}
        
        .blur-fade {{
            position: absolute;
            bottom: 0;
            left: 0;
            right: 0;
            height: 200px;
            background: linear-gradient(to bottom, transparent, white 80%);
            z-index: 10;
        }}
        
        .paywall {{
            position: absolute;
            bottom: 50px;
            left: 50%;
            transform: translateX(-50%);
            text-align: center;
            z-index: 20;
            background: white;
            padding: 30px 50px;
            border-radius: 12px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.15);
            border: 2px solid #667eea;
        }}
        
        .paywall-icon {{
            font-size: 48px;
            margin-bottom: 15px;
        }}
        
        .paywall h3 {{
            font-size: 20px;
            margin-bottom: 10px;
            color: #333;
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        }}
        
        .paywall p {{
            color: #666;
            margin-bottom: 20px;
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            font-size: 14px;
        }}
        
        .paywall-price {{
            font-size: 32px;
            font-weight: bold;
            color: #667eea;
            margin-bottom: 20px;
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        }}
        
        .paywall-btn {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            padding: 15px 40px;
            font-size: 16px;
            border-radius: 8px;
            cursor: pointer;
            transition: transform 0.2s, box-shadow 0.2s;
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            font-weight: 600;
        }}
        
        .paywall-btn:hover {{
            transform: translateY(-2px);
            box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
        }}
        
        .watermark {{
            position: fixed;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%) rotate(-45deg);
            font-size: 80px;
            color: rgba(200, 200, 200, 0.15);
            pointer-events: none;
            z-index: 100;
            font-weight: bold;
            text-transform: uppercase;
            letter-spacing: 10px;
        }}
        
        .preview-badge {{
            position: fixed;
            top: 20px;
            right: 20px;
            background: #ff6b6b;
            color: white;
            padding: 8px 16px;
            border-radius: 20px;
            font-size: 12px;
            font-weight: bold;
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            z-index: 1000;
        }}
        
        @media print {{
            .paywall, .blur-fade, .preview-badge {{
                display: none !important;
            }}
            .blur-content {{
                filter: none !important;
                opacity: 1 !important;
            }}
        }}
    </style>
</head>
<body>
    <div class="preview-badge">👁 ПРЕДПРОСМОТР</div>
    
    <div class="watermark">ПРЕВЬЮ</div>
    
    <div class="preview-container">
        <div class="document-header">
            <div class="document-title">{doc_type}</div>
            <div class="document-type">Предварительный просмотр документа</div>
        </div>
        
        <div class="content">
            {formatted_content}
        </div>
        
        <div class="blur-overlay">
            <div class="blur-content">
                <p style="text-indent: 20px;">
                    Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.
                </p>
                <p style="text-indent: 20px;">
                    Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.
                </p>
                <p style="text-indent: 20px;">
                    Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium doloremque laudantium, totam rem aperiam, eaque ipsa quae ab illo inventore veritatis et quasi architecto beatae vitae dicta sunt explicabo.
                </p>
                <p style="text-indent: 20px;">
                    Nemo enim ipsam voluptatem quia voluptas sit aspernatur aut odit aut fugit, sed quia consequuntur magni dolores eos qui ratione voluptatem sequi nesciunt.
                </p>
            </div>
            <div class="blur-fade"></div>
            
            <div class="paywall">
                <div class="paywall-icon">🔒</div>
                <h3>Продолжение документа скрыто</h3>
                <p>Оплатите документ, чтобы получить полную версию<br>в форматах PDF и DOCX</p>
                <div class="paywall-price">{case_data.get('price', 5000)} ₽</div>
                <button class="paywall-btn" onclick="window.parent.postMessage('pay', '*')">
                    💳 Оплатить и скачать
                </button>
            </div>
        </div>
    </div>
</body>
</html>'''
    
    return html


def format_content_to_html(content):
    """Convert plain text to HTML with formatting."""
    if not content:
        return ""
    
    lines = content.strip().split('\n')
    html_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        upper_line = line.upper()
        
        # Main title
        if any(word in upper_line for word in ['ИСКОВОЕ ЗАЯВЛЕНИЕ', 'АПЕЛЛЯЦИОННАЯ ЖАЛОБА', 
                                                  'ПРЕТЕНЗИЯ', 'СТРАТЕГИЯ ЗАЩИТЫ', 'ВОЗРАЖЕНИЕ']):
            html_lines.append(f'<h1>{escape_html(line)}</h1>')
        # Numbered sections
        elif re.match(r'^\d+\.', line):
            html_lines.append(f'<h2>{escape_html(line)}</h2>')
        # Special sections
        elif any(line.startswith(word) for word in ['ПРОШУ:', 'ТРЕБУЕМ:', 'ПРИЛОЖЕНИЕ:', 'ПРИЛОЖЕНИЯ:']):
            html_lines.append(f'<h2>{escape_html(line)}</h2>')
        # Party info
        elif any(word in line for word in ['Истец:', 'Ответчик:', 'Апеллянт:', 'Отправитель:', 'Получатель:']):
            html_lines.append(f'<p style="font-weight: bold;">{escape_html(line)}</p>')
        # Regular paragraph
        else:
            html_lines.append(f'<p>{escape_html(line)}</p>')
    
    return '\n'.join(html_lines)


def escape_html(text):
    """Escape HTML special characters."""
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&#x27;'))


def generate_preview_docx(full_content, output_path, visible_ratio=0.25):
    """
    Generate DOCX preview with only visible portion.
    
    Args:
        full_content: Full document content
        output_path: Path to save preview
        visible_ratio: How much content to include
    
    Returns:
        Path to preview file
    """
    from document_generator import set_document_style, add_heading_custom, add_paragraph_custom
    
    visible_content, _ = split_content_for_preview(full_content, visible_ratio)
    
    doc = Document()
    set_document_style(doc)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add preview watermark header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('👁 ПРЕДВАРИТЕЛЬНЫЙ ПРОСМОТР')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 107, 107)
    
    doc.add_paragraph()  # Spacing
    
    # Add visible content
    lines = visible_content.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        upper_line = line.upper()
        
        if any(word in upper_line for word in ['ИСКОВОЕ ЗАЯВЛЕНИЕ', 'АПЕЛЛЯЦИОННАЯ ЖАЛОБА']):
            add_heading_custom(doc, line, level=1)
        elif re.match(r'^\d+\.', line):
            add_heading_custom(doc, line, level=2)
        elif any(line.startswith(word) for word in ['ПРОШУ:', 'ТРЕБУЕМ:', 'ПРИЛОЖЕНИЕ:']):
            add_heading_custom(doc, line, level=2)
        elif any(word in line for word in ['Истец:', 'Ответчик:']):
            add_paragraph_custom(doc, line, bold=True)
        else:
            add_paragraph_custom(doc, line)
    
    # Add "Continue" message
    doc.add_paragraph()
    msg = doc.add_paragraph()
    msg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = msg.add_run('\n\n[Продолжение документа доступно после оплаты]')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(102, 126, 234)
    
    # Add pay info
    doc.add_paragraph()
    pay = doc.add_paragraph()
    pay.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pay.add_run('Оплатите документ, чтобы получить полную версию в форматах PDF и DOCX')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.save(output_path)
    return output_path
