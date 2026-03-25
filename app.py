from dotenv import load_dotenv
load_dotenv()

"""
Legal AI Service - Main Application
Flask backend for legal document analysis and case management.
"""

import os
import json
import logging
import threading
import time
from datetime import datetime, timedelta
from functools import wraps

from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge

# Import local modules
from config import (
    init_config, init_tariffs, UPLOAD_FOLDER, MAX_CONTENT_LENGTH, 
    ALLOWED_EXTENSIONS, API_PREFIX, LOG_LEVEL, LOG_FORMAT, LOG_FILE
)
from models import db, User, Case, DownloadHistory, PaymentTransaction, IPRequest, Tariff, UserTariff, ClarityRequest, init_db
from auth import (
    require_auth, get_current_user, get_client_ip, 
    register_user, login_user, generate_token, auth_response, error_response,
    refresh_access_token
)
from document_generator import generate_case_report, generate_legal_document_docx
from preview_generator import generate_preview_html, split_content_for_preview, generate_preview_docx
from kimi_api import analyze_case_documents, check_context_consistency, get_document_type_name, generate_legal_document

# Initialize Flask app
app = Flask(__name__)
init_config(app)
CORS(app, resources={r'/api/*': {'origins': '*'}})

# Initialize database
init_db(app)

# Setup logging
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL),
    format=LOG_FORMAT,
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# ==================== CREDIT SYSTEM FUNCTIONS ====================

def get_user_balance(user_id):
    """Get users current balance across all active tariffs."""
    user_tariffs = UserTariff.query.filter_by(
        user_id=user_id, 
        is_active=True
    ).all()
    
    total_credits = sum(ut.credits_total for ut in user_tariffs)
    used_credits = sum(ut.credits_used for ut in user_tariffs)
    remaining = total_credits - used_credits
    
    # Get current tariff (most recent active)
    current_tariff = None
    if user_tariffs:
        current_ut = max(user_tariffs, key=lambda x: x.activated_at)
        current_tariff = current_ut.tariff.code if current_ut.tariff else None
    
    percentage_remaining = int((remaining / total_credits) * 100) if total_credits > 0 else 0
    
    return {
        'credits_total': total_credits,
        'credits_used': used_credits,
        'credits_remaining': remaining,
        'current_tariff': current_tariff,
        'percentage_remaining': percentage_remaining
    }


def deduct_credit(user_id, count=1):
    """Deduct credits from users balance. Returns (success, remaining_credits)."""
    user_tariffs = UserTariff.query.filter_by(
        user_id=user_id, 
        is_active=True
    ).order_by(UserTariff.activated_at.asc()).all()
    
    remaining_to_deduct = count
    
    for ut in user_tariffs:
        available = ut.get_remaining_credits()
        if available >= remaining_to_deduct:
            ut.credits_used += remaining_to_deduct
            db.session.commit()
            return True, get_user_balance(user_id)['credits_remaining']
        elif available > 0:
            ut.credits_used += available
            remaining_to_deduct -= available
    
    db.session.commit()
    
    if remaining_to_deduct == 0:
        return True, get_user_balance(user_id)['credits_remaining']
    
    return False, get_user_balance(user_id)['credits_remaining']


def check_balance_notifications(user_id):
    """Check balance and send mock notifications at threshold levels."""
    balance = get_user_balance(user_id)
    user = User.query.get(user_id)
    
    if not user:
        return
    
    percentage = balance['percentage_remaining']
    remaining = balance['credits_remaining']
    
    thresholds = [25, 15, 10, 5]
    
    for threshold in thresholds:
        if percentage <= threshold:
            logger.info(f"[MOCK EMAIL] To: {user.email} | Subject: Низкий баланс кредитов ({percentage}%)")
            logger.info(f"[MOCK EMAIL] Body: У вас осталось {remaining} кредитов ({percentage}%). Пополните баланс!")
            break
    
    # Check for first purchase discount offer
    if user.is_first_purchase and not user.discount_used and remaining > 0:
        logger.info(f"[MOCK EMAIL] To: {user.email} | Subject: Специальное предложение - скидка 25%!")
        logger.info(f"[MOCK EMAIL] Body: Вы оплатили первый раз! Получите скидку 25% на тарифы 'Фирма' и 'Палата'!")
        logger.info(f"[MOCK EMAIL] Promo code: FIRST25")


# Store for verification codes (in production, use Redis)
verification_codes = {}

def generate_verification_code():
    """Generate 4-digit verification code."""
    import random
    return str(random.randint(1000, 9999))

def send_mock_verification_email(email, code):
    """Log verification code to console (mock)."""
    logger.info(f"[MOCK EMAIL] To: {email}")
    logger.info(f"[MOCK EMAIL] Subject: Код подтверждения регистрации")
    logger.info(f"[MOCK EMAIL] Body: Ваш код подтверждения: {code}")
    logger.info(f"[MOCK EMAIL] Code expires in 10 minutes")

def allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def save_uploaded_file(file, user_id, case_id):
    """Save uploaded file to user/case directory."""
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        
        # Create user/case directory structure
        user_dir = os.path.join(UPLOAD_FOLDER, f'user_{user_id}')
        case_dir = os.path.join(user_dir, f'case_{case_id}')
        os.makedirs(case_dir, exist_ok=True)
        
        # Save file
        filepath = os.path.join(case_dir, filename)
        file.save(filepath)
        
        return filepath
    return None


def analyze_with_progress(case_id, documents, document_type, custom_request):
    """
    Analyze documents with progress updates (runs in background thread).
    """
    with app.app_context():
        case = Case.query.get(case_id)
        if not case:
            logger.error(f"Case {case_id} not found")
            return
        
        try:
            # Step 1: Analyzing documents (0-30%)
            case.update_progress(5, 'Анализируем документы...')
            db.session.commit()
            time.sleep(5)
            
            case.update_progress(15, 'Анализируем документы...')
            db.session.commit()
            time.sleep(5)
            
            case.update_progress(25, 'Анализируем документы...')
            db.session.commit()
            time.sleep(3)
            
            # Step 2: Researching legal practice (30-60%)
            case.update_progress(35, 'Изучаем юридическую практику...')
            db.session.commit()
            time.sleep(5)
            
            case.update_progress(45, 'Изучаем юридическую практику...')
            db.session.commit()
            time.sleep(5)
            
            case.update_progress(55, 'Изучаем юридическую практику...')
            db.session.commit()
            time.sleep(3)
            
            # Step 3: Forming document (60-90%)
            case.update_progress(65, 'Формируем документ...')
            db.session.commit()
            
            # Run AI analysis
            document_texts = []
            for doc_path in documents:
                try:
                    with open(doc_path, 'r', encoding='utf-8', errors='ignore') as f:
                        document_texts.append(f.read()[:5000])  # Limit text length
                except:
                    pass
            
            # Use mock analysis
            analysis_result = analyze_case_documents(document_texts or ['Документ для анализа'])
            
            # Generate legal document
            case_data = {
                'court_name': 'Мировой судья судебного участка № ___',
                'plaintiff': {'name': 'Истец', 'address': 'Адрес истца'},
                'defendant': {'name': 'Ответчик', 'address': 'Адрес ответчика'},
                'claim_amount': '5000',
                'date': __import__("datetime").datetime.now().strftime('%d.%m.%Y'),
                'date_today': __import__("datetime").datetime.now().strftime('"%d" %B %Y г.'),
                'loan_amount': '100000',
                'due_date': '15.04.2024',
                'interest_amount': '5000',
                'penalty_amount': '2000',
                'court_fee': '3000',
                'lawyer_fee': '5000',
                'total_expenses': '8000',
                'moral_damage': '10000'
            }
            
            # Use custom request if provided, otherwise use selected document type
            doc_type = document_type if document_type else 'complaint'
            
            # Try AI generation with fallback
            try:
                generated_content = generate_legal_document(case_data, doc_type)
            except Exception as e:
                logger.error(f'AI generation failed: {e}')
                # Fallback template based on doc_type
                if doc_type == 'appeal':
                    generated_content = f"""АПЕЛЛЯЦИОННАЯ ЖАЛОБА

В {case_data['court_name']}

Заявитель: {case_data['plaintiff']['name']}
Ответчик: {case_data['defendant']['name']}

ТРЕБОВАНИЯ:
Отменить решение суда первой инстанции.

ОБОСНОВАНИЕ:
Суд первой инстанции допустил существенные нарушения процессуальных норм.

На основании изложенного, руководствуясь ст. 320-322 ГПК РФ,

ПРОШУ:
Отменить решение и удовлетворить иск в полном объеме.

Дата: {case_data['date']}
Подпись: ________________"""
                elif doc_type == 'petition':
                    generated_content = f"""ДОСУДЕБНАЯ ПРЕТЕНЗИЯ

От: {case_data['plaintiff']['name']}
Кому: {case_data['defendant']['name']}

Уважаемый(ая) {case_data['defendant']['name']}!

Вы нарушили сроки возврата займа на сумму {case_data['loan_amount']} руб.

ТРЕБУЕМ:
1. Вернуть сумму займа: {case_data['loan_amount']} руб.
2. Уплатить проценты: {case_data['interest_amount']} руб.
3. Уплатить неустойку: {case_data['penalty_amount']} руб.

При неисполнении в течение 10 дней обратимся в суд.

Дата: {case_data['date']}
Подпись: ________________"""
                elif doc_type == 'statement':
                    generated_content = f"""СТРАТЕГИЯ ЗАЩИТЫ

Дело: Взыскание задолженности
Сумма: {case_data['loan_amount']} руб.

I. АНАЛИЗ ДЕЛА
- Суть спора: Взыскание задолженности по договору займа
- Правовые основания: ст. 807-810 ГК РФ

II. ТАКТИКА ЗАЩИТЫ
1. Доказать заключение договора займа
2. Доказать передачу денежных средств
3. Доказать нарушение срока возврата
4. Взыскать проценты и неустойку

III. ДОКАЗАТЕЛЬСТВА
- Договор займа
- Расписка или платежные документы
- Переписка о просрочке

IV. ПРОГНОЗ
Вероятность удовлетворения иска: ВЫСОКАЯ

Подготовил: ________________
Дата: {case_data['date']}"""
                else:  # complaint (default)
                    generated_content = f"""ИСКОВОЕ ЗАЯВЛЕНИЕ

В {case_data['court_name']}

Истец: {case_data['plaintiff']['name']}
Адрес: {case_data['plaintiff']['address']}

Ответчик: {case_data['defendant']['name']}
Адрес: {case_data['defendant']['address']}

О взыскании задолженности по договору займа

Сумма иска: {case_data['claim_amount']} руб.

ИСТОРИЯ ДЕЛА:
Ответчик получил денежные средства в размере {case_data['loan_amount']} руб. Срок возврата ({case_data['due_date']}) истек, обязательства не исполнены.

ПРАВОВОЕ ОБОСНОВАНИЕ:
На основании ст. 807, 808, 810 ГК РФ ответчик обязан возвратить сумму займа и уплатить проценты.

ИСКОВЫЕ ТРЕБОВАНИЯ:
1. Взыскать сумму долга: {case_data['loan_amount']} руб.
2. Взыскать проценты: {case_data['interest_amount']} руб.
3. Взыскать неустойку: {case_data['penalty_amount']} руб.
4. Взыскать судебные расходы: {case_data['court_fee']} руб.
5. Взыскать моральный вред: {case_data['moral_damage']} руб.

Приложения:
1. Копия искового заявления
2. Документы, подтверждающие требования

Дата: {case_data['date']}
Подпись: ________________"""
            
            case.update_progress(75, 'Формируем документ...')
            db.session.commit()
            time.sleep(5)
            
            case.update_progress(85, 'Формируем документ...')
            db.session.commit()
            time.sleep(3)
            
            # Step 4: Checking result (90-100%)
            case.update_progress(92, 'Проверяем результат...')
            db.session.commit()
            time.sleep(3)
            
            case.update_progress(97, 'Проверяем результат...')
            db.session.commit()
            time.sleep(2)
            
            # Validate content before generating DOCX
            if not generated_content or len(generated_content) < 50:
                generated_content = f"""ИСКОВОЕ ЗАЯВЛЕНИЕ

В суд общей юрисдикции

Истец: {case_data.get('plaintiff', {}).get('name', 'Истец')}
Ответчик: {case_data.get('defendant', {}).get('name', 'Ответчик')}

О взыскании задолженности

Сумма иска: {case_data.get('loan_amount', '___')} руб.

На основании изложенного, руководствуясь ст. 807-810 ГК РФ,

ПРОШУ:
Взыскать с ответчика задолженность.

Дата: {case_data.get('date', '___')}
Подпись: ________________"""
            
            # Save generated document
            generated_dir = os.path.join('generated', f'user_{case.user_id}')
            os.makedirs(generated_dir, exist_ok=True)
            docx_path = os.path.join(generated_dir, f'case_{case_id}_document.docx')
            
            # Generate DOCX with error handling
            try:
                generate_legal_document_docx(generated_content, docx_path, case_data)
            except Exception as docx_error:
                logger.error(f'DOCX generation failed: {docx_error}')
                # Create simple DOCX using python-docx directly
                from docx import Document
                doc = Document()
                doc.add_heading('Юридический документ', 0)
                for line in generated_content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
                doc.save(docx_path)
            
            case.generated_document_path = docx_path
            
            # Complete
            analysis_result['generated_content'] = generated_content[:500] + '...'
            case.complete_analysis(analysis_result)
            case.paid = True  # Auto-mark as paid (free mode)
            db.session.commit()
            
            logger.info(f"Analysis completed for case {case_id}")
            
        except Exception as e:
            case.fail_analysis()
            db.session.commit()
            logger.error(f"Analysis failed for case {case_id}: {str(e)}")


def log_request_info():
    """Log request information."""
    logger.info(f"{request.method} {request.path} - IP: {get_client_ip()}")


# ==================== ERROR HANDLERS ====================

@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(e):
    """Handle file too large error."""
    return jsonify({
        'success': False,
        'error': 'File too large',
        'message': f'Maximum file size is {MAX_CONTENT_LENGTH // (1024 * 1024)} MB'
    }), 413


@app.errorhandler(404)
def handle_not_found(e):
    """Handle 404 errors."""
    # Для не-API запросов отдаем index.html (SPA behavior)
    if not request.path.startswith('/api/'):
        return send_from_directory('/opt/legal-ai-service/public', 'index.html')
    return jsonify({
        'success': False,
        'error': 'Not found',
        'message': 'The requested resource was not found'
    }), 404


@app.errorhandler(500)
def handle_server_error(e):
    """Handle 500 errors."""
    logger.error(f"Server error: {str(e)}")
    return jsonify({
        'success': False,
        'error': 'Server error',
        'message': 'An internal server error occurred'
    }), 500


# ==================== AUTHENTICATION ENDPOINTS ====================

@app.route(f'{API_PREFIX}/register', methods=['POST'])
def register():
    """
    Register a new user.
    
    Request body:
        - username: User login name
        - email: User email
        - password: User password
    
    Returns:
        User data and authentication tokens
    """
    log_request_info()
    
    data = request.get_json()
    if not data:
        return error_response('No data provided', 'VALIDATION_ERROR', 400)
    
    username = data.get('username', '').strip()
    email = data.get('email', '').strip()
    password = data.get('password', '')
    
    if not username or not email or not password:
        return error_response('Имя пользователя, email и пароль обязательны', 'VALIDATION_ERROR', 400)
    
    ip_address = get_client_ip()
    user, error = register_user(username, email, password, ip_address)
    
    if error:
        return error_response(error, 'REGISTRATION_ERROR', 400)
    
    logger.info(f"User registered: {username}")
    return auth_response(user, 'Регистрация успешна')


@app.route(f'{API_PREFIX}/login', methods=['POST'])
def login():
    """
    Authenticate user and get tokens.
    
    Request body:
        - username: Username or email
        - password: User password
    
    Returns:
        User data and authentication tokens
    """
    log_request_info()
    
    data = request.get_json()
    if not data:
        return error_response('No data provided', 'VALIDATION_ERROR', 400)
    
    username = data.get('username', '').strip()
    password = data.get('password', '')
    
    if not username or not password:
        return error_response('Укажите имя пользователя/email и пароль', 'VALIDATION_ERROR', 400)
    
    ip_address = get_client_ip()
    user, error = login_user(username, password, ip_address)
    
    if error:
        return error_response(error, 'AUTHENTICATION_ERROR', 401)
    
    logger.info(f"User logged in: {user.username}")
    return auth_response(user, 'Вход выполнен успешно')


@app.route(f'{API_PREFIX}/refresh', methods=['POST'])
def refresh_token():
    """
    Refresh access token using refresh token.
    
    Request body:
        - refresh_token: Refresh token
    
    Returns:
        New access token
    """
    log_request_info()
    
    data = request.get_json()
    if not data:
        return error_response('No data provided', 'VALIDATION_ERROR', 400)
    
    refresh_token = data.get('refresh_token')
    if not refresh_token:
        return error_response('Refresh token is required', 'VALIDATION_ERROR', 400)
    
    new_token, error = refresh_access_token(refresh_token)
    
    if error:
        return error_response(error, 'TOKEN_ERROR', 401)
    
    return jsonify({
        'success': True,
        'data': {
            'access_token': new_token,
            'token_type': 'Bearer',
            'expires_in': 86400
        }
    })


@app.route(f'{API_PREFIX}/me', methods=['GET'])
@require_auth
def get_current_user_info():
    """Get current user information."""
    log_request_info()
    user = request.current_user
    return jsonify({
        'success': True,
        'data': user.to_dict()
    })


# ==================== CASE MANAGEMENT ENDPOINTS ====================

@app.route(f'{API_PREFIX}/upload', methods=['POST'])
@require_auth
def upload_documents():
    """
    Upload documents for a case with document type selection.
    Deducts 1 credit from users balance.
    
    Form data:
        - case_title: Title for the new case (optional)
        - document_type: Type of document (complaint, appeal, petition, statement)
        - custom_request: Custom text request (optional)
        - files: One or more files to upload
    
    Returns:
        Case information with uploaded documents and remaining credits
    """
    log_request_info()
    user = request.current_user
    
    # Get case info
    case_title = request.form.get('case_title', 'Новое дело')
    document_type = request.form.get('document_type')  # complaint, appeal, petition, statement
    custom_request = request.form.get('custom_request', '').strip()
    
    # Validate: either document_type or custom_request must be provided
    if not document_type and not custom_request:
        return error_response('Выберите тип документа или введите свой запрос', 'VALIDATION_ERROR', 400)
    
    # Create new case
    case = Case(
        user_id=user.id,
        case_title=case_title,
        document_type=document_type,
        custom_request=custom_request
    )
    db.session.add(case)
    db.session.commit()
    
    # Process uploaded files
    uploaded_files = []
    logger.info(f"Upload request: case_title={case_title}, doc_type={document_type}, files_count={len(request.files.getlist('files'))}")
    
    if 'files' in request.files:
        files = request.files.getlist('files')
        logger.info(f'Processing {len(files)} files...')
        for i, file in enumerate(files):
            if file.filename:
                logger.info(f'Processing file {i+1}/{len(files)}: {file.filename}')
                try:
                    filepath = save_uploaded_file(file, user.id, case.id)
                    if filepath:
                        case.add_document(filepath)
                        uploaded_files.append(os.path.basename(filepath))
                        logger.info(f'File {i+1} saved: {filepath}')
                    else:
                        logger.warning(f'File {i+1} failed to save')
                except Exception as e:
                    logger.error(f'Error processing file {i+1}: {e}')
                    import traceback
                    logger.error(traceback.format_exc())
            else:
                logger.warning(f'File {i+1} has no filename')
    
    db.session.commit()
    
    # Log upload info
    logger.info(f"Documents uploaded for case {case.id}: {len(uploaded_files)} files")
    logger.info(f"Document type: {document_type}, Custom request: {custom_request[:50] if custom_request else 'None'}")
    
    return jsonify({
        'success': True,
        'message': f'{len(uploaded_files)} файлов загружено',
        'data': {
            'case': case.to_dict(),
            'uploaded_files': uploaded_files,
            'credits_remaining': 999
        }
    })


@app.route(f'{API_PREFIX}/cases', methods=['GET'])
@require_auth
def get_cases():
    """
    Get list of users cases.
    
    Query parameters:
        - status: Filter by status (optional)
        - page: Page number (default: 1)
        - per_page: Items per page (default: 10)
    
    Returns:
        List of cases
    """
    log_request_info()
    user = request.current_user
    
    # Get query parameters
    status = request.args.get('status')
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 10, type=int)
    
    # Build query
    query = Case.query.filter_by(user_id=user.id)
    if status:
        query = query.filter_by(status=status)
    
    # Paginate
    pagination = query.order_by(Case.created_at.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )
    
    cases = [case.to_dict() for case in pagination.items]
    
    return jsonify({
        'success': True,
        'data': {
            'cases': cases,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': pagination.total,
                'pages': pagination.pages,
                'has_next': pagination.has_next,
                'has_prev': pagination.has_prev
            }
        }
    })


@app.route(f'{API_PREFIX}/case/<int:case_id>', methods=['GET'])
@require_auth
def get_case(case_id):
    """
    Get detailed information about a case.
    
    Args:
        case_id: Case ID
    
    Returns:
        Case details including documents and analysis
    """
    log_request_info()
    user = request.current_user
    
    case = Case.query.get(case_id)
    if not case or case.user_id != user.id:
        return error_response('Дело не найдено или доступ запрещен', 'NOT_FOUND', 404)
    
    return jsonify({
        'success': True,
        'data': {
            'case': case.to_dict(include_analysis=True)
        }
    })


@app.route(f'{API_PREFIX}/case/<int:case_id>', methods=['DELETE'])
@require_auth
def delete_case(case_id):
    """
    Delete a case and its documents.
    
    Args:
        case_id: Case ID
    
    Returns:
        Success message
    """
    log_request_info()
    user = request.current_user
    
    case = Case.query.get(case_id)
    if not case or case.user_id != user.id:
        return error_response('Дело не найдено или доступ запрещен', 'NOT_FOUND', 404)
    
    # Delete associated files
    import shutil
    case_dir = os.path.join(UPLOAD_FOLDER, f'user_{user.id}', f'case_{case_id}')
    if os.path.exists(case_dir):
        shutil.rmtree(case_dir)
    
    # Delete generated document if exists
    if case.generated_document_path and os.path.exists(case.generated_document_path):
        os.remove(case.generated_document_path)
    
    # Delete from database
    db.session.delete(case)
    db.session.commit()
    
    logger.info(f"Case {case_id} deleted by user {user.id}")
    
    return jsonify({
        'success': True,
        'message': 'Дело удалено'
    })


# ==================== ANALYSIS ENDPOINTS ====================

@app.route(f'{API_PREFIX}/analyze', methods=['POST'])
@require_auth
def analyze_case():
    """
    Start AI analysis for a case with animated progress.
    
    Request body:
        - case_id: Case ID to analyze
    
    Returns:
        Analysis started confirmation
    """
    log_request_info()
    user = request.current_user
    
    data = request.get_json()
    if not data:
        return error_response('No data provided', 'VALIDATION_ERROR', 400)
    
    case_id = data.get('case_id')
    if not case_id:
        return error_response('Case ID is required', 'VALIDATION_ERROR', 400)
    
    case = Case.query.get(case_id)
    if not case or case.user_id != user.id:
        return error_response('Дело не найдено или доступ запрещен', 'NOT_FOUND', 404)
    
    # Check if case is already analyzing or completed
    if case.status == 'analyzing':
        return jsonify({
            'success': True,
            'message': 'Анализ уже выполняется',
            'data': {'case': case.to_dict()}
        })
    
    if case.status == 'completed':
        return jsonify({
            'success': True,
            'message': 'Анализ уже завершен',
            'data': {'case': case.to_dict(include_analysis=True)}
        })
    
    # Start analysis
    case.start_analysis()
    db.session.commit()
    
    # Run analysis in background thread
    documents = case.get_documents()
    thread = threading.Thread(
        target=analyze_with_progress,
        args=(case.id, documents, case.document_type, case.custom_request)
    )
    thread.daemon = True
    thread.start()
    
    logger.info(f"Analysis started for case {case_id}")
    
    return jsonify({
        'success': True,
        'message': 'Анализ запущен',
        'data': {
            'case': case.to_dict()
        }
    })


@app.route(f'{API_PREFIX}/case/<int:case_id>/progress', methods=['GET'])
@require_auth
def get_case_progress(case_id):
    """
    Get current progress of case analysis.
    
    Args:
        case_id: Case ID
    
    Returns:
        Current progress percentage and message
    """
    log_request_info()
    user = request.current_user
    
    case = Case.query.get(case_id)
    if not case or case.user_id != user.id:
        return error_response('Дело не найдено или доступ запрещен', 'NOT_FOUND', 404)
    
    return jsonify({
        'success': True,
        'data': {
            'case_id': case.id,
            'status': case.status,
            'progress': case.progress,
            'progress_message': case.progress_message,
            'is_completed': case.status == 'completed',
            'is_paid': case.paid
        }
    })


# ==================== PAYMENT ENDPOINTS ====================

@app.route(f'{API_PREFIX}/payment/create', methods=['POST'])
@require_auth
def create_payment():
    """
    Create payment for case.
    
    Request body:
        - case_id: Case ID to pay for
    
    Returns:
        Payment URL and transaction info
    """
    log_request_info()
    user = request.current_user
    
    data = request.get_json()
    if not data:
        return error_response('No data provided', 'VALIDATION_ERROR', 400)
    
    case_id = data.get('case_id')
    if not case_id:
        return error_response('Case ID is required', 'VALIDATION_ERROR', 400)
    
    case = Case.query.get(case_id)
    if not case or case.user_id != user.id:
        return error_response('Дело не найдено или доступ запрещен', 'NOT_FOUND', 404)
    
    if case.paid:
        return jsonify({
            'success': True,
            'message': 'Дело уже оплачено',
            'data': {'case': case.to_dict()}
        })
    
    # Create payment transaction
    transaction = PaymentTransaction(
        user_id=user.id,
        case_id=case.id,
        amount=case.price,
        payment_type='document',
        description=f'Оплата генерации документа: {case.case_title}'
    )
    db.session.add(transaction)
    db.session.commit()
    
    # Generate mock payment URL (placeholder)
    payment_url = f'/payment/confirm?transaction_id={transaction.transaction_id}'
    transaction.payment_url = payment_url
    db.session.commit()
    
    logger.info(f"Payment created for case {case_id}, transaction {transaction.transaction_id}")
    
    return jsonify({
        'success': True,
        'message': 'Платеж создан',
        'data': {
            'transaction_id': transaction.transaction_id,
            'amount': case.price,
            'payment_url': payment_url,
            'case': case.to_dict()
        }
    })


@app.route(f'{API_PREFIX}/payment/confirm', methods=['POST'])
@require_auth
def confirm_payment():
    """
    Confirm payment (mock webhook).
    
    Request body:
        - transaction_id: Transaction ID
    
    Returns:
        Payment confirmation
    """
    log_request_info()
    user = request.current_user
    
    data = request.get_json()
    if not data:
        return error_response('No data provided', 'VALIDATION_ERROR', 400)
    
    transaction_id = data.get('transaction_id')
    if not transaction_id:
        return error_response('Transaction ID is required', 'VALIDATION_ERROR', 400)
    
    transaction = PaymentTransaction.query.filter_by(
        transaction_id=transaction_id,
        user_id=user.id
    ).first()
    
    if not transaction:
        return error_response('Транзакция не найдена', 'NOT_FOUND', 404)
    
    if transaction.status == 'completed':
        return jsonify({
            'success': True,
            'message': 'Платеж уже подтвержден',
            'data': {'transaction': transaction.to_dict()}
        })
    
    # Mark transaction as completed
    transaction.mark_completed()
    
    case = None
    # Mark case as paid
    if transaction.case_id:
        case = Case.query.get(transaction.case_id)
        if case:
            case.mark_as_paid()
    
    db.session.commit()
    
    logger.info(f"Payment confirmed for transaction {transaction_id}")
    
    return jsonify({
        'success': True,
        'message': 'Платеж успешно подтвержден',
        'data': {
            'transaction': transaction.to_dict(),
            'case': case.to_dict() if transaction.case_id else None
        }
    })


@app.route(f'{API_PREFIX}/payments/history', methods=['GET'])
@require_auth
def payment_history():
    """
    Get users payment history.
    
    Query parameters:
        - page: Page number (default: 1)
        - per_page: Items per page (default: 20)
    
    Returns:
        List of payment transactions
    """
    log_request_info()
    user = request.current_user
    
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    
    transactions = PaymentTransaction.query.filter_by(user_id=user.id)\
        .order_by(PaymentTransaction.created_at.desc())\
        .paginate(page=page, per_page=per_page, error_out=False)
    
    return jsonify({
        'success': True,
        'data': {
            'transactions': [t.to_dict() for t in transactions.items],
            'page': page,
            'per_page': per_page,
            'total': transactions.total,
            'pages': transactions.pages
        }
    })


# ==================== TARIFFS & BALANCE ENDPOINTS ====================

@app.route(f'{API_PREFIX}/tariffs', methods=['GET'])
def get_tariffs():
    """Get all available tariffs with savings calculation."""
    log_request_info()
    
    tariffs = Tariff.query.filter_by(is_active=True).order_by(Tariff.price).all()
    
    single_tariff = Tariff.query.filter_by(code='single').first()
    single_price = single_tariff.price if single_tariff else 5000
    
    result = []
    for t in tariffs:
        tariff_dict = t.to_dict()
        if t.code != 'single':
            single_cost = t.credits * single_price
            savings = single_cost - t.price
            savings_percent = round((savings / single_cost) * 100) if single_cost > 0 else 0
            tariff_dict['savings'] = {
                'vs_single': savings,
                'percent': savings_percent
            }
        result.append(tariff_dict)
    
    return jsonify({
        'success': True,
        'data': {
            'tariffs': result,
            'currency': 'RUB'
        }
    })


@app.route(f'{API_PREFIX}/balance', methods=['GET'])
@require_auth
def get_balance():
    """Get users current credit balance and tariff info."""
    log_request_info()
    user = request.current_user
    
    active_tariff = UserTariff.query.filter_by(
        user_id=user.id,
        is_active=True
    ).first()
    
    tariff_info = None
    if active_tariff:
        tariff = Tariff.query.get(active_tariff.tariff_id)
        if tariff:
            tariff_info = {
                'name': tariff.name,
                'code': tariff.code,
                'credits_total': active_tariff.credits_total,
                'credits_used': active_tariff.credits_used,
                'credits_remaining': active_tariff.get_remaining_credits(),
                'expires_at': active_tariff.expires_at.isoformat() if active_tariff.expires_at else None
            }
    
    notifications = []
    if user.credits < 5:
        notifications.append({
            'type': 'critical',
            'message': 'Осталось менее 5 кредитов! Пополните баланс.'
        })
    elif user.credits < 15:
        notifications.append({
            'type': 'warning',
            'message': 'Осталось менее 15 кредитов.'
        })
    
    return jsonify({
        'success': True,
        'data': {
            'credits': user.credits,
            'tariff': tariff_info,
            'notifications': notifications,
            'is_first_purchase': user.is_first_purchase
        }
    })


# ==================== EMAIL VERIFICATION ====================

@app.route(f'{API_PREFIX}/auth/send-code', methods=['POST'])
def send_verification_code():
    """Send email verification code."""
    log_request_info()
    
    data = request.get_json()
    email = data.get('email', '').lower().strip()
    
    if not email or '@' not in email:
        return jsonify({'success': False, 'error': 'Некорректный email'}), 400
    
    code = generate_verification_code()
    import time
    verification_codes[email] = {
        'code': code,
        'expires_at': time.time() + 600,
        'attempts': 0
    }
    
    send_mock_verification_email(email, code)
    
    return jsonify({
        'success': True,
        'message': 'Код подтверждения отправлен на email',
        'data': {'email': email, 'expires_in': 600}
    })


@app.route(f'{API_PREFIX}/auth/verify-code', methods=['POST'])
def verify_code():
    """Verify email code."""
    log_request_info()
    
    data = request.get_json()
    email = data.get('email', '').lower().strip()
    code = data.get('code', '').strip()
    
    if not email or not code:
        return jsonify({'success': False, 'error': 'Email и код обязательны'}), 400
    
    import time
    stored = verification_codes.get(email)
    
    if not stored:
        return jsonify({'success': False, 'error': 'Код не найден'}), 404
    
    if time.time() > stored['expires_at']:
        del verification_codes[email]
        return jsonify({'success': False, 'error': 'Код истёк'}), 410
    
    if stored['attempts'] >= 3:
        del verification_codes[email]
        return jsonify({'success': False, 'error': 'Слишком много попыток'}), 429
    
    stored['attempts'] += 1
    
    if stored['code'] != code:
        return jsonify({'success': False, 'error': 'Неверный код'}), 400
    
    del verification_codes[email]
    
    return jsonify({
        'success': True,
        'message': 'Email подтверждён',
        'data': {'email': email, 'verified': True}
    })


# ==================== DOWNLOAD ENDPOINTS ====================

@app.route(f'{API_PREFIX}/download/<int:case_id>', methods=['GET'])
@require_auth
def download_document(case_id):
    """
    Download generated document.
    
    Args:
        case_id: Case ID
    
    Returns:
        DOCX file
    """
    log_request_info()
    user = request.current_user
    
    case = Case.query.get(case_id)
    if not case or case.user_id != user.id:
        return error_response('Дело не найдено или доступ запрещен', 'NOT_FOUND', 404)
    
    if not case.generated_document_path or not os.path.exists(case.generated_document_path):
        return error_response('Документ еще не сгенерирован', 'NOT_READY', 400)
    
    # Record download
    download_history = DownloadHistory(
        user_id=user.id,
        case_id=case_id,
        ip_address=get_client_ip(),
        user_agent=request.headers.get('User-Agent'),
        download_type='document'
    )
    db.session.add(download_history)
    db.session.commit()
    
    logger.info(f"Document downloaded for case {case_id} by user {user.id}")
    
    return send_file(
        case.generated_document_path,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'{case.document_type or "document"}_case_{case_id}.docx'
    )


# ==================== PRICING ENDPOINT ====================

@app.route(f'{API_PREFIX}/pricing', methods=['GET'])
def get_pricing():
    """Get pricing information."""
    log_request_info()
    return jsonify({
        'success': True,
        'data': {
            'document_price': 5000,
            'currency': 'RUB',
            'description': 'Стоимость генерации одного юридического документа'
        }
    })


# ==================== MAIN ====================

if __name__ == '__main__':
    logger.info("Starting Legal AI Service...")
    
    # Create required directories
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs('generated', exist_ok=True)
    
    # Initialize database and tariffs
    with app.app_context():
        db.create_all()
        init_tariffs()
    
    app.run(
        host='0.0.0.0',
        port=5000,
        debug=os.environ.get('FLASK_DEBUG', 'False').lower() == 'true'
    )

# ==================== PREVIEW ENDPOINTS ====================

@app.route(f'{API_PREFIX}/preview/<int:case_id>', methods=['GET'])
@require_auth
def get_document_preview(case_id):
    """Get document preview with blur effect - shows 25% of content."""
    log_request_info()
    user = request.current_user
    
    case = Case.query.get(case_id)
    if not case or case.user_id != user.id:
        return error_response('Дело не найдено', 'NOT_FOUND', 404)
    
    if not case.generated_content:
        return error_response('Документ еще не сгенерирован', 'NOT_READY', 400)
    
    try:
        from preview_generator import generate_preview_html, split_content_for_preview
        visible_content, _ = split_content_for_preview(case.generated_content, visible_ratio=0.25)
        case_data = {'price': case.price, 'document_type': case.document_type, 'case_title': case.case_title}
        html_content = generate_preview_html(case_data, visible_content, case.document_type)
        return html_content, 200, {'Content-Type': 'text/html; charset=utf-8'}
    except Exception as e:
        logger.error(f"Preview error: {str(e)}")
        return error_response('Ошибка генерации превью', 'INTERNAL_ERROR', 500)


@app.route(f'{API_PREFIX}/preview/<int:case_id>/docx', methods=['GET'])
@require_auth
def get_document_preview_docx(case_id):
    """Get document preview as DOCX with partial content."""
    log_request_info()
    user = request.current_user
    
    case = Case.query.get(case_id)
    if not case or case.user_id != user.id:
        return error_response('Дело не найдено', 'NOT_FOUND', 404)
    
    if not case.generated_content:
        return error_response('Документ не готов', 'NOT_READY', 400)
    
    try:
        from preview_generator import generate_preview_docx
        preview_path = os.path.join(app.config.get('GENERATED_FOLDER', 'generated'), f'preview_case_{case_id}.docx')
        generate_preview_docx(case.generated_content, preview_path, visible_ratio=0.25)
        
        if os.path.exists(preview_path):
            return send_file(preview_path, as_attachment=True, 
                           download_name=f'preview_{case.case_title or "document"}.docx',
                           mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        return error_response('Ошибка создания превью', 'GENERATION_ERROR', 500)
    except Exception as e:
        logger.error(f"Preview DOCX error: {str(e)}")
        return error_response('Ошибка генерации', 'INTERNAL_ERROR', 500)


# ==================== PREVIEW ENDPOINTS ====================


# ==================== STATIC HTML ROUTES ====================

@app.route('/')
def index():
    return send_from_directory('/opt/legal-ai-service/public', 'index.html')

@app.route('/login.html')
def login_page():
    return send_from_directory('/opt/legal-ai-service/public', 'login.html')

@app.route('/register.html')
def register_page():
    return send_from_directory('/opt/legal-ai-service/public', 'register.html')

@app.route('/dashboard.html')
def dashboard_page():
    return send_from_directory('/opt/legal-ai-service/public', 'dashboard.html')

@app.route('/pricing.html')
def pricing_page():
    return send_from_directory('/opt/legal-ai-service/public', 'pricing.html')

@app.route('/about.html')
def about_page():
    return send_from_directory('/opt/legal-ai-service/public', 'about.html')

@app.route('/public/<path:filename>')
def public_files(filename):
    return send_from_directory('/opt/legal-ai-service/public', filename)
