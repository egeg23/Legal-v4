"""
Document generator module for Legal AI Service.
Generates DOCX documents from legal content.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_document_style(doc):
    """Set default document styles."""
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set paragraph format
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(6)


def add_heading_custom(doc, text, level=1):
    """Add a custom heading."""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading


def add_paragraph_custom(doc, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a custom paragraph."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    paragraph.alignment = alignment
    return paragraph


def generate_legal_document_docx(content, output_path, case_data=None):
    """
    Generate a legal document as DOCX file.
    
    Args:
        content: Document content (text)
        output_path: Path to save the generated DOCX
        case_data: Optional case data for template variables
    
    Returns:
        Path to generated file
    """
    doc = Document()
    
    # Set document styles
    set_document_style(doc)
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Process content line by line
    lines = content.strip().split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Detect headings and special formatting
        upper_line = line.upper()
        
        if 'ИСКОВОЕ ЗАЯВЛЕНИЕ' in upper_line or 'АПЕЛЛЯЦИОННАЯ ЖАЛОБА' in upper_line or \
           'ПРЕТЕНЗИЯ' in upper_line or 'СТРАТЕГИЯ ЗАЩИТЫ' in upper_line:
            # Main document title
            add_heading_custom(doc, line, level=1)
        elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or \
             line.startswith('4.') or line.startswith('5.') or line.startswith('6.') or \
             line.startswith('7.') or line.startswith('8.') or line.startswith('9.'):
            # Section heading
            add_heading_custom(doc, line, level=2)
        elif line.startswith('ПРОШУ:') or line.startswith('ТРЕБУЕМ:') or \
             line.startswith('ПРИЛОЖЕНИЕ:') or line.startswith('ПРИЛОЖЕНИЯ:'):
            # Special sections
            add_heading_custom(doc, line, level=2)
        elif line.startswith('В ') and 'суд' in line.lower():
            # Court name - center align
            add_paragraph_custom(doc, line, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        elif 'Истец:' in line or 'Ответчик:' in line or 'Апеллянт:' in line or \
             'Отправитель:' in line or 'Получатель:' in line:
            # Party information
            add_paragraph_custom(doc, line, bold=True)
        else:
            # Regular paragraph
            add_paragraph_custom(doc, line)
    
    # Add signature section
    doc.add_paragraph()
    signature_para = doc.add_paragraph()
    signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature_run = signature_para.add_run(
        f"\n«___» ___________ {datetime.now().year} г.\n\n_________________ / _________________ /"
    )
    signature_run.font.name = 'Times New Roman'
    signature_run.font.size = Pt(12)
    signature_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    # Save document
    doc.save(output_path)
    
    return output_path


def generate_case_report(case_data, output_dir):
    """
    Generate a case analysis report as DOCX.
    
    Args:
        case_data: Case data dictionary
        output_dir: Directory to save the report
    
    Returns:
        Path to generated report
    """
    doc = Document()
    set_document_style(doc)
    
    # Title
    title = doc.add_heading('АНАЛИТИЧЕСКИЙ ОТЧЕТ ПО ДЕЛУ', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Case info
    doc.add_heading('1. ИНФОРМАЦИЯ О ДЕЛЕ', level=2)
    add_paragraph_custom(doc, f"Название дела: {case_data.get('case_title', 'Не указано')}")
    add_paragraph_custom(doc, f"Дата создания: {case_data.get('created_at', 'Не указано')}")
    add_paragraph_custom(doc, f"Статус: {case_data.get('status', 'Не указано')}")
    
    # Analysis results
    analysis = case_data.get('analysis_result', {})
    
    if analysis:
        doc.add_heading('2. РЕЗУЛЬТАТЫ АНАЛИЗА', level=2)
        
        # Key findings
        if 'key_findings' in analysis:
            doc.add_heading('2.1. Ключевые выводы', level=3)
            for finding in analysis['key_findings']:
                add_paragraph_custom(doc, f"• {finding}")
        
        # Legal issues
        if 'legal_issues' in analysis:
            doc.add_heading('2.2. Правовые вопросы', level=3)
            for issue in analysis['legal_issues']:
                add_paragraph_custom(doc, f"• {issue}")
        
        # Risks
        if 'risks' in analysis:
            doc.add_heading('2.3. Риски', level=3)
            for risk in analysis['risks']:
                add_paragraph_custom(doc, f"• {risk.get('description', '')} (Уровень: {risk.get('level', '')})")
                if 'mitigation' in risk:
                    add_paragraph_custom(doc, f"  Рекомендация: {risk['mitigation']}")
        
        # Recommendations
        if 'recommendations' in analysis:
            doc.add_heading('2.4. Рекомендации', level=3)
            for rec in analysis['recommendations']:
                add_paragraph_custom(doc, f"• {rec}")
    
    # Generated content
    if 'generated_content' in analysis:
        doc.add_heading('3. СГЕНЕРИРОВАННЫЙ ДОКУМЕНТ', level=2)
        add_paragraph_custom(doc, analysis['generated_content'])
    
    # Save report
    report_path = os.path.join(output_dir, f'case_{case_data.get("id", "unknown")}_report.docx')
    doc.save(report_path)
    
    return report_path


